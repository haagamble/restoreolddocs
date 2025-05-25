# This program can restore a word document that used TransCyrillicSans font for Tajik.
# The Tajik text sometimes appear as Wingdings or other characters instead of Cyrillic.
# This script replaces the TransCyrillicSans font text with the correct Tajik characters
# using predefined mappings for single characters and double-character pairs.
# Usage: python restore.py
# The text to be restored should be TransCyrillicSans so the script can identify it.
# Any text in other fonts will be left unchanged.

from docx import Document
from subs import char_map, pairs_map

def replace_text_in_run(run, char_map, pairs_map):
    font_name = run.font.name
    if font_name and "TransCyrillicSans" in font_name:
        text = run.text
        # Replace double-character pairs first
        for pair, tajik_char in pairs_map.items():
            text = text.replace(pair, tajik_char)
        # Replace remaining single characters
        text = ''.join(char_map.get(char, char) for char in text)
        run.text = text
        # Optionally, reset font to a Unicode font so it displays properly:
        run.font.name = "Arial"

def fix_text_in_tables(doc, char_map, pairs_map):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        replace_text_in_run(run, char_map, pairs_map)

def fix_text_in_paragraphs(doc, char_map, pairs_map):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            replace_text_in_run(run, char_map, pairs_map)

if __name__ == "__main__":
    # Load the document. 
    # Assumes the document is in docs folder and named fixme.docx
    # Change the path and name of output file as needed.
    doc = Document("docs/fixme.docx")
    fix_text_in_paragraphs(doc, char_map, pairs_map)
    fix_text_in_tables(doc, char_map, pairs_map)
    doc.save("docs_fixed/fixed.docx")
    print("Done fixing document")
